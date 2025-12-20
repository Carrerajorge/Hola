from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .schemas import DocSpec, HeadingBlock, ParagraphBlock, BulletsBlock, TableBlock, PageBreakBlock


def _set_default_styles(doc: Document) -> None:
    # Sensible defaults. You can extend this with a full design system.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_toc(doc: Document) -> None:
    # Word will populate the TOC when the user updates fields (or on open depending on settings).
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.addnext(fld)
    doc.add_paragraph()  # spacing after TOC


def _add_table(doc: Document, block: TableBlock) -> None:
    n_cols = len(block.columns)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = block.style or "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(block.columns):
        run = hdr_cells[j].paragraphs[0].add_run(str(col))
        run.bold = True

    for row in block.rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = "" if val is None else str(val)

    doc.add_paragraph()  # spacing after table


def render_word(spec: DocSpec, out_path: str) -> str:
    doc = Document()
    _set_default_styles(doc)

    # Title
    if spec.title:
        doc.add_heading(spec.title, level=0)

    if spec.add_toc:
        _add_toc(doc)

    # Content blocks
    for b in spec.blocks:
        if isinstance(b, HeadingBlock):
            doc.add_heading(b.text, level=b.level)
        elif isinstance(b, ParagraphBlock):
            doc.add_paragraph(b.text)
        elif isinstance(b, BulletsBlock):
            for item in b.items:
                doc.add_paragraph(item, style="List Bullet")
        elif isinstance(b, TableBlock):
            _add_table(doc, b)
        elif isinstance(b, PageBreakBlock):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            # If you add new block types, handle them here.
            doc.add_paragraph(str(getattr(b, "text", "")))

    if spec.author:
        doc.core_properties.author = spec.author

    doc.save(out_path)
    return out_path
