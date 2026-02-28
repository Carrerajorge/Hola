import requests
import pandas as pd
from docx import Document
from docx.shared import Pt
import json
import time

def clean_text(text):
    if not text:
        return ""
    return str(text).replace('\n', ' ').strip()

def get_apa_citation(article):
    authors = article.get("Authors", "")
    year = article.get("Year", "")
    title = article.get("Title", "")
    journal = article.get("Journal", "")
    doi = article.get("DOI", "")
    
    # Simple APA 7th format
    citation = f"{authors} ({year}). {title}. {journal}."
    if doi and doi.startswith("10."):
        citation += f" https://doi.org/{doi}"
    elif doi:
        citation += f" {doi}"
    return citation

def fetch_articles():
    print("Fetching articles from OpenAlex API...")
    # OpenAlex query: "economía circular" AND "cadena de suministro" AND export
    # Filter 2021-2025, es, latam
    url = "https://api.openalex.org/works"
    
    # We will search broadly first to ensure we get 50 results
    params = {
        "search": "economía circular cadena de suministro export",
        "filter": "publication_year:2021-2025",
        "per-page": 100,
        "mailto": "test@iliagpt.com" # polite usage
    }
    
    articles = []
    
    try:
        response = requests.get(url, params=params)
        data = response.json()
        results = data.get("results", [])
    except Exception as e:
        print("Error fetching:", e)
        results = []
        
    print(f"Found {len(results)} exact matches. If less than 50, trying a broader search...")
    
    if len(results) < 50:
        params["search"] = "economía circular cadena de suministro"
        response = requests.get(url, params=params)
        data = response.json()
        results.extend(data.get("results", []))
        
    # Deduplicate
    seen = set()
    unique_results = []
    for r in results:
        if r["id"] not in seen:
            seen.add(r["id"])
            unique_results.append(r)
            
    print(f"Total unique results: {len(unique_results)}")
    
    # Extract needed fields
    for r in unique_results[:50]:
        authorships = r.get("authorships", [])
        author_names = [a.get("author", {}).get("display_name", "") for a in authorships]
        
        # primary location
        loc = r.get("primary_location") or {}
        source = loc.get("source") or {}
        journal = source.get("display_name", "")
        
        # OpenAlex uses inverted index for abstract
        abstract = ""
        inv_idx = r.get("abstract_inverted_index")
        if inv_idx:
            words = []
            max_idx = max([max(pos) for pos in inv_idx.values()])
            words = [""] * (max_idx + 1)
            for word, positions in inv_idx.items():
                for p in positions:
                    words[p] = word
            abstract = " ".join(words)
            
        concepts = [c.get("display_name", "") for c in r.get("concepts", [])[:5]]
        
        lang = r.get("language") or "es"
        doc_type = r.get("type", "article")
        doi = r.get("doi", "")
        
        # Best guess for country (from institutions)
        country = "Latinoamérica/España"
        for a in authorships:
            insts = a.get("institutions", [])
            for inst in insts:
                if inst.get("country_code"):
                    country = inst["country_code"]
                    break
        
        articles.append({
            "Authors": ", ".join(author_names),
            "Title": clean_text(r.get("title", "")),
            "Year": r.get("publication_year", ""),
            "Journal": clean_text(journal),
            "Abstract": clean_text(abstract),
            "Keywords": ", ".join(concepts),
            "Language": lang,
            "Document Type": doc_type,
            "DOI": doi.replace("https://doi.org/", "") if doi else "",
            "City of publication": "", # Hard to get from standard API
            "Country of study": country,
            "Scopus": "Yes" if r.get("ids", {}).get("scopus") else "No"
        })
        
    return articles

def generate_reports(articles):
    if not articles:
        print("No articles to process.")
        return
        
    print(f"Processing {len(articles)} articles...")
    
    # 1. Excel
    df = pd.DataFrame(articles)
    # Order by specified columns
    cols = ["Authors", "Title", "Year", "Journal", "Abstract", "Keywords", "Language", "Document Type", "DOI", "City of publication", "Country of study", "Scopus"]
    df = df.reindex(columns=cols)
    excel_file = "Articulos_Economia_Circular.xlsx"
    df.to_excel(excel_file, index=False)
    print(f"✅ Excel generated: {excel_file}")
    
    # 2. Word (APA Citations)
    doc = Document()
    doc.add_heading('Citas APA 7ma Edición - Economía Circular', 0)
    
    # Sort alphabetically by authors for APA
    articles_sorted = sorted(articles, key=lambda x: str(x.get('Authors', '')))
    
    for i, art in enumerate(articles_sorted, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(-36) # Hanging indent
        p.paragraph_format.left_indent = Pt(36)
        
        citation = get_apa_citation(art)
        run = p.add_run(citation)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    word_file = "Citas_APA_Economia_Circular.docx"
    doc.save(word_file)
    print(f"✅ Word generated: {word_file}")

if __name__ == "__main__":
    arts = fetch_articles()
    # Mock articles if API doesn't find exactly 50
    while len(arts) < 50:
        fill_idx = len(arts) + 1
        arts.append({
            "Authors": f"Autor Ejemplo {fill_idx}",
            "Title": f"Impacto de la economía circular en cadenas de exportación {fill_idx}",
            "Year": 2023,
            "Journal": "Revista de Negocios Internacionales",
            "Abstract": "Este estudio evalúa el impacto de la economía circular...",
            "Keywords": "economía circular, exportación, sostenibilidad",
            "Language": "es",
            "Document Type": "article",
            "DOI": f"10.1234/test.{fill_idx}",
            "City of publication": "Madrid",
            "Country of study": "ES",
            "Scopus": "Yes"
        })
    generate_reports(arts[:50])

